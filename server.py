# server.py
from fastmcp import FastMCP
from docx import Document
from docx.shared import Inches
import os

mcp = FastMCP("Demo 🚀")

@mcp.tool(
    name="summary",
    description="总结会话历史记录，支持总结当前会话或全部会话历史，调用本工具时直接返回结果"
)
def summary(summaryHistory: str, historyType: str = "all") -> str:
    """summary history"""
    prefix = "全部会话历史记录：" if historyType == "all" else "当前会话历史记录："
    return prefix + summaryHistory

@mcp.tool(
    name="export_to_word",
    description="将会话历史记录导出到Word文档"
)
def export_to_word(summaryHistory: str, historyType: str = "all", filePath: str = r"C:\Users\25137\Desktop\每日日报.docx") -> str:
    """export history to Word document"""
    # 创建Word文档
    doc = Document()
    
    # 添加标题
    title = "全部会话历史记录" if historyType == "all" else "当前会话历史记录"
    doc.add_heading(title, 0)
    
    # 添加内容
    doc.add_paragraph(summaryHistory)
    
    # 保存文件
    # 确保目录存在
    directory = os.path.dirname(filePath)
    if directory and not os.path.exists(directory):
        os.makedirs(directory)
        
    doc.save(filePath)
    
    return f"已成功导出到Word文档：{filePath}"

if __name__ == "__main__":
    mcp.run()